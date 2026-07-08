"""
canopy.core.transcriber — YouTube subtitle fetch and transcript conversion.

Fetches VTT subtitles from yt-dlp info_dict URLs, strips timestamps/HTML,
and saves the result as plain .txt or Word .docx.
"""

import os
import re
import urllib.request

try:
    from docx import Document as _DocxDoc
    from docx.shared import Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def best_lang(info: dict) -> tuple[str, str]:
    """Return (lang_code, display_label) for the best available subtitle.

    Prefers English manual captions, then English auto-generated, then any
    manual language, then any auto-generated language.
    Returns ('', '') if no subtitles exist at all.
    """
    subs = info.get("subtitles", {})
    auto = info.get("automatic_captions", {})
    for source, kind in [(subs, "manual captions"), (auto, "auto-generated")]:
        if "en" in source:
            return "en", f"English · {kind}"
        for lang in source:
            if lang.startswith("en"):
                return lang, f"English · {kind}"
    for source, kind in [(subs, "manual captions"), (auto, "auto-generated")]:
        for lang in source:
            return lang, f"{lang} · {kind}"
    return "", ""


def _vtt_url(info: dict, lang: str) -> str | None:
    """Return the VTT subtitle URL for `lang` (manual first, auto second)."""
    for source in (info.get("subtitles", {}), info.get("automatic_captions", {})):
        entries = source.get(lang, [])
        for e in entries:
            if e.get("ext") == "vtt":
                return e.get("url")
        if entries:
            return entries[0].get("url")
    return None


def fetch_vtt(info: dict, lang: str) -> str:
    """Download and return raw VTT text for the given language code."""
    url = _vtt_url(info, lang)
    if not url:
        raise RuntimeError(
            f"No subtitles found for language '{lang}'. "
            "This video may not have captions."
        )
    with urllib.request.urlopen(url, timeout=30) as resp:
        return resp.read().decode("utf-8", errors="replace")


def vtt_to_text(vtt: str) -> str:
    """Parse VTT subtitle content into clean, readable plain text.

    YouTube auto-generated captions use a rolling 2-line window: each cue
    carries forward the tail of the previous cue as context, producing a
    pattern like:

        "...My name"                          ← partial A
        "...My name is Marshall Cruz..."      ← A + B (longer)
        "is Marshall Cruz..."                 ← partial B
        "is Marshall Cruz...of DMs..."        ← B + C (longer)

    We fix this with three passes:
      1. Drop any block that is a strict prefix of the next block (it's just
         an incomplete version that gets completed in the following cue).
      2. Strip the prefix overlap each block shares with the previous kept
         block, leaving only genuinely new content.
      3. Join fragments into sentence-boundary paragraphs so the output reads
         naturally rather than breaking mid-phrase.
    """
    # ── Pass 0: parse raw cue text blocks ────────────────────────────────────
    blocks: list[str] = []
    current: list[str] = []
    past_header = False

    for raw in vtt.splitlines():
        line = raw.strip()

        if not past_header:
            if not line:
                past_header = True
            continue

        if re.match(r"[\d:]+\.?\d*\s+-->", line):
            if current:
                blocks.append(" ".join(current))
                current = []
            continue

        if not line:
            if current:
                blocks.append(" ".join(current))
                current = []
            continue

        if re.match(r"^\d+$", line):
            continue

        clean = re.sub(r"<[^>]+>", "", line)
        for old, new in (("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">"), ("&nbsp;", " ")):
            clean = clean.replace(old, new)
        clean = clean.strip()
        if clean:
            current.append(clean)

    if current:
        blocks.append(" ".join(current))

    # ── Pass 1: drop blocks that are strict prefixes of the next block ───────
    filtered: list[str] = []
    for i, block in enumerate(blocks):
        nxt = blocks[i + 1] if i + 1 < len(blocks) else ""
        if nxt.startswith(block) and nxt != block:
            continue
        filtered.append(block)

    # ── Pass 2: strip prefix overlap shared with the previous kept block ─────
    fragments: list[str] = []
    for i, block in enumerate(filtered):
        if i == 0:
            fragments.append(block)
            continue
        prev = fragments[-1]
        new_part = block
        # Find the longest suffix of prev that is also a prefix of block;
        # only strip if the overlap is >= 5 chars (avoids false positives).
        for length in range(min(len(prev), len(block)), 4, -1):
            if prev.endswith(block[:length]):
                candidate = block[length:].lstrip()
                if candidate:
                    new_part = candidate
                break
        if new_part and new_part != prev:
            fragments.append(new_part)

    # ── Pass 3: join fragments into sentence-boundary paragraphs ─────────────
    _SENT_END = re.compile(r"[.?!]\s*$")
    paras: list[str] = []
    run: list[str] = []

    for frag in fragments:
        run.append(frag)
        if _SENT_END.search(frag):
            paras.append(" ".join(run))
            run = []

    if run:
        paras.append(" ".join(run))

    return "\n".join(paras)


def _safe_filename(title: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', "_", title)[:80].strip()


def save_txt(text: str, title: str, save_dir: str) -> str:
    """Save transcript as UTF-8 plain text. Returns the saved file path."""
    path = os.path.join(save_dir, f"{_safe_filename(title)} — Transcript.txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(title + "\n")
        f.write("=" * min(len(title), 80) + "\n\n")
        f.write(text)
    return path


def save_docx(text: str, title: str, uploader: str,
              duration: str, source_url: str, save_dir: str) -> str:
    """Save transcript as Word .docx. Returns the saved file path."""
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. "
            "Run: pip install python-docx --break-system-packages"
        )

    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    heading = doc.add_heading(title, 0)
    if heading.runs:
        heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    meta = "  ·  ".join(p for p in (uploader, duration) if p)
    if meta:
        p = doc.add_paragraph(meta)
        if p.runs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    if source_url:
        p = doc.add_paragraph(source_url)
        if p.runs:
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x33, 0x88, 0x44)

    doc.add_paragraph()

    for line in text.splitlines():
        line = line.strip()
        if line:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    path = os.path.join(save_dir, f"{_safe_filename(title)} — Transcript.docx")
    doc.save(path)
    return path
